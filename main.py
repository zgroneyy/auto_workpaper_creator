#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jan 24 2022 02:33:58 AM
@author: ozgur oney
"""
import os
import random
    
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from tkinter import *
from tkinter import messagebox
from tkinter.filedialog import askopenfilename
from datetime import date

window = Tk()
window.title('Workpaper Creator')
# set minimum&maximum window size value - for fixed window
window.minsize(600, 400)
window.maxsize(600, 400)
window.config(bg='#456')
f = ('sans-serif', 13)
btn_font = ('sans-serif', 10)
bgcolor = '#BF5517'

#array keeping unique title values
genvar = StringVar()
genopt = ['Denetçi Yardımcısı', 
          'BT Denetçi Yardımcısı', 
          'Yetkili Denetçi Yardımcısı', 
          'Yetkili BT Denetçi Yardımcısı', 
          'Denetçi', 
          'Müfettiş', 
          'Başdenetçi', 
          'Başmüfettiş']
genvar.set('BT Denetçi Yardımcısı')

#function of Clear button
def clear_inputs():
    baslangictarihi.delete(0, 'end')
    bitistarihi.delete(0, 'end')
    denetim_elemani_adi.delete(0, 'end')
    pathLabel.delete(0,'end')
#change user-input text of date to timestamp object
def date_to_timestamp(d) :
    day, month, year = d.split('/')
    return date(int(year), int(month), int(day))
#create random date in given interval
def randomDate(start, end):
    stime = date_to_timestamp(start)
    etime = date_to_timestamp(end)
    ptime = stime + random.random() * (etime-stime)
    return str(ptime)
#select file contains audit steps
def select_file():
    #hide window (optional)
    window.withdraw()
    #define action
    file_path = askopenfilename(title="Open file", 
                                       filetypes=[("Word",".docx"),("TXT",".txt"),
                                                  ("All files",".*")])
    #if selected path is valid
    if file_path != "":
        print ("you chose file with path:", file_path)

    else:
        print ("you didn't open anything!")
    #add selected path to label, so that user can see it
    pathLabel.delete(0, END)
    pathLabel.insert(0, file_path)
    #mutate path so that its valid in all OS'
    file_path = os.path.dirname(file_path)
    #bring back the lights! you can see window after selection
    window.deiconify()
    return file_path

def generate():
    #get file path from label
    file_path = pathLabel.get()
    
    with open (file_path, encoding='utf8') as f:
        # read lines for auto-numbering
        lines = f.readlines()
        lines = [line.rstrip() for line in lines]
        #for each line in the doc
        for line in lines:
            #create document
            document = Document()
            # Create a character level style Object ("CommentsStyle") then defines its parameters
            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentStyle', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(12)
            obj_font.name = 'Times New Roman'
            #add header in bold
            baslik = document.add_paragraph()
            baslik.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            baslik.add_run('S-100 ÇALIŞMA KAĞIDI', style='CommentStyle').bold=True
            #create a random date, add it to document
            tarih = document.add_paragraph()
            tarih.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            tarih.add_run('Tarih: ', style='CommentStyle').bold=True
            tarih.add_run(randomDate(baslangictarihi.get(), bitistarihi.get()), style='CommentStyle')
            #add workpaper number in increasing order for each step
            ck_num= document.add_paragraph()
            ck_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ck_num.add_run('Çalışma Kâğıdı Numarası: ', style='CommentStyle').bold=True
            index = [x for x in range(len(lines)) if line in lines[x]]
            ck_num.add_run(str(index[0]+1).strip("[]"), style='CommentStyle')
            #add audit step
            denetim_adimi = document.add_paragraph()
            denetim_adimi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            denetim_adimi.add_run('İlgili Denetim Adımı: ', style='CommentStyle').bold=True
            denetim_adimi.add_run(line, style='CommentStyle')
            #add auditor name /with his/her title
            denetim_elemani = document.add_paragraph()
            denetim_elemani.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            denetim_elemani.add_run('Testi Gerçekleştiren Denetim Elemanı:  ', style='CommentStyle').bold=True
            isim = str(denetim_elemani_adi.get()) +  ", " +  str(genvar.get()) + ", Denetim Genel Müdürlüğü"
            denetim_elemani.add_run(isim, style='CommentStyle')
            #sampling 
            orneklem = document.add_paragraph()
            orneklem.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            orneklem.add_run('Örneklem Yöntemi:  ', style='CommentStyle').bold=True
            #documents viewed/researched
            incelenen = document.add_paragraph()
            incelenen.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            incelenen.add_run('İncelenen Dokümanlar:  ', style='CommentStyle').bold=True
            #related (if exists) non-conformity number
            bulgu_num = document.add_paragraph()
            bulgu_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bulgu_num.add_run('Bulgu Numarası:  ', style='CommentStyle').bold=True
            #test procedure
            test_prod = document.add_paragraph()
            test_prod.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            test_prod.add_run('Ayrıntılı Test Prosedürü:  ', style='CommentStyle').bold=True
            #detect filename
            filename = line.rstrip() + '.docx'
            #save document
            document.save(filename)     
    #message informing last user about process
    messagebox.showinfo("Sonuc", "Dosya(lar) başarıyla oluşturuldu.")
    #void
    return None
# frame carrying items in the window
frame = Frame(window, padx=20, pady=20, bg=bgcolor)
frame.pack(expand=True, fill=BOTH)
#labels
Label(frame, text= "Den. Programı: ", font=f, 
      bg=bgcolor).grid(column=0, row=0, padx=15, pady=15)
pathLabel = Entry(frame,textvariable="")
pathLabel.grid(column=1, row=0, padx=15, pady=15)
Label(frame, text = "Baş. tarihi : ", font=f, 
      bg=bgcolor).grid(column=0, row=1, padx=15, pady=15)
btn_frame = Frame(frame, bg=bgcolor)
btn_frame.grid(columnspan=2, pady=(50, 0))
Label(frame, text = "Bit. tarihi : ", font=f, 
      bg=bgcolor).grid(column=0, row=2, padx=15, pady=15)
btn_frame = Frame(frame, bg=bgcolor)
btn_frame.grid(columnspan=2, pady=(50, 0))
Label(frame, text = "Den. Elemanı İsim: ", font=f, 
      bg=bgcolor).grid(column=0, row=3, padx=15, pady=15)
btn_frame = Frame(frame, bg=bgcolor)
btn_frame.grid(columnspan=2, pady=(50, 0))
Label(frame, text = "Title: ", font=f, 
      bg=bgcolor).grid(column=0, row=4, padx=15, pady=15)
btn_frame = Frame(frame, bg=bgcolor)
btn_frame.grid(columnspan=2, pady=(50, 0))
#input widgets   
fileSelect = Button(
    frame,
    text='Dosya Seç',
    command=select_file,
    font=btn_font,
    padx=10, 
    pady=5,
    width=3
)
fileSelect.grid(row=0, column=2)
baslangictarihi = Entry(frame, width=20, font=f)
baslangictarihi.grid(row=1, column=1)
bitistarihi = Entry(frame, width=20, font=f)
bitistarihi.grid(row=2, column=1)
denetim_elemani_adi = Entry(frame, width=20, font=f)
denetim_elemani_adi.grid(row=3, column=1)
rutbe = OptionMenu(
    frame, 
    genvar,
    *genopt
)
rutbe.grid(row=4, column=1, pady=(5,0))
rutbe.config(width=16, font=f)

#default values
baslangictarihi.insert(0,'01/01/2022')
bitistarihi.insert(0,'14/04/2022')
denetim_elemani_adi.insert(0,'John Doe')

#buttons
submit_btn = Button(
    btn_frame,
    text='Oluştur',
    command=generate,
    font=btn_font,
    padx=10, 
    pady=5
)
submit_btn.pack(side=LEFT, expand=True, padx=(15, 0))

clear_btn = Button(
    btn_frame,
    text='Temizle',
    command=clear_inputs,
    font=btn_font,
    padx=10, 
    pady=5,
    width=7
)
clear_btn.pack(side=LEFT, expand=True, padx=15)

exit_btn = Button(
    btn_frame,
    text='Çıkış',
    command=lambda:window.destroy(),
    font=btn_font,
    padx=10, 
    pady=5
)
exit_btn.pack(side=LEFT, expand=True)

# mainloop
window.mainloop()