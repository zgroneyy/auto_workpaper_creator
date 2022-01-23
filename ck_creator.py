#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Jan 11 22:32:24 2022

@author: ozgur
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import *
from docx.enum.style import *
from docx.shared import Pt

def __init__(self):
    self.auditor_name = ''

def setAuditor(self, auditor_name):
    self.auditor_name = auditorname

def getAuditor(self):
    return self.auditor_name

def creator():    
    with open ('list.txt', encoding='utf8') as f:
        # read lines for auto-numbering
        lines = f.readlines()
        lines = [line.rstrip() for line in lines]
        # for each line in the doc
        for line in lines:
            
            document = Document()
            # Creates a character level style Object ("CommentsStyle") then defines its parameters
            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentStyle', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(12)
            obj_font.name = 'Times New Roman'
            
            baslik = document.add_paragraph()
            baslik.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            baslik_run = baslik.add_run('S-100 ÇALIŞMA KAĞIDI', style='CommentStyle').bold=True
            
            tarih = document.add_paragraph()
            tarih.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            u_tarih = input("Tarih giriniz: ")
            tarih_run = tarih.add_run('Tarih: ', style='CommentStyle').bold=True
            u_tarih = tarih.add_run(u_tarih, style='CommentStyle')
            
            ck_num= document.add_paragraph()
            ck_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ck_num.add_run('Çalışma Kâğıdı Numarası: ', style='CommentStyle').bold=True
            index = [x for x in range(len(lines)) if line in lines[x]]
            ck_num.add_run(str(index[0]+1).strip("[]"), style='CommentStyle')
            
            denetim_adimi = document.add_paragraph()
            denetim_adimi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            denetim_adimi.add_run('İlgili Denetim Adımı: ', style='CommentStyle').bold=True
            denetim_adimi.add_run(line, style='CommentStyle')
            
            denetim_elemani = document.add_paragraph()
            denetim_elemani.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            denetim_elemani.add_run('Testi Gerçekleştiren Denetim Elemanı:  ', style='CommentStyle').bold=True
            denetim_elemani.add_run('Özgür Öney, Denetim Genel Müdürlüğü, Yetkili BT Denetçi Yardımcısı', style='CommentStyle')
            # print(gui.getAuditorName())
            # denetim_elemani.add_run(gui.getAuditorName(), style='CommentStyle')
            
            orneklem = document.add_paragraph()
            orneklem.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            orneklem.add_run('Örneklem Yöntemi:  ', style='CommentStyle').bold=True
            
            incelenen = document.add_paragraph()
            incelenen.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            incelenen.add_run('İncelenen Dokümanlar:  ', style='CommentStyle').bold=True
    
            bulgu_num = document.add_paragraph()
            bulgu_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bulgu_num.add_run('Bulgu Numarası:  ', style='CommentStyle').bold=True
    
            test_prod = document.add_paragraph()
            test_prod.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            test_prod.add_run('Ayrıntılı Test Prosedürü:  ', style='CommentStyle').bold=True
    
            filename = line.rstrip() + '.docx'
            document.save(filename)